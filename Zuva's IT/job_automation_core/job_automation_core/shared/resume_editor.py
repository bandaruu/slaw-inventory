#!/usr/bin/env python3
"""
Surgical Resume Editor
Edits DOCX resumes with targeted modifications instead of full rewrites
Preserves original formatting and authentic voice
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Optional
import re
from copy import deepcopy


class SurgicalResumeEditor:
    """
    Edit DOCX resumes surgically - make targeted changes while preserving formatting.
    This approach maintains authenticity vs full AI rewrites.
    """
    
    def __init__(self, base_resume_path: str):
        """
        Load base resume for editing.
        
        Args:
            base_resume_path: Path to original .docx resume
        """
        self.doc = Document(base_resume_path)
        self.base_path = base_resume_path
        self.modifications = []
    
    def inject_keywords(self, keywords: List[str], skills: List[str]) -> int:
        """
        Layer 1: Inject keywords into existing content.
        Swaps synonyms with exact JD keywords where natural.

        Args:
            keywords: Keywords from job description
            skills: Technical skills from job description

        Returns:
            Number of keywords injected
        """
        injected_count = 0

        # Build synonym map for intelligent replacement
        synonym_map = self._build_synonym_map(keywords)

        # Inject into paragraphs
        for paragraph in self.doc.paragraphs:
            original_text = paragraph.text

            # Replace synonyms with JD keywords, operating on individual runs
            for jd_keyword, synonyms in synonym_map.items():
                for synonym in synonyms:
                    pattern = r'\b' + re.escape(synonym) + r'\b'
                    if re.search(pattern, original_text, re.IGNORECASE):
                        replaced = self._replace_in_runs(paragraph, pattern, jd_keyword)
                        if replaced:
                            injected_count += 1
                            self.modifications.append({
                                'type': 'keyword_injection',
                                'original': original_text,
                                'modified': paragraph.text
                            })
                        break

        # Add missing skills to Skills section
        skills_added = self._add_skills_to_section(skills)
        injected_count += skills_added

        return injected_count
    
    def rewrite_summary(self, new_summary: str) -> bool:
        """
        Layer 2: Replace Professional Summary section.
        This is the ONLY section that gets fully rewritten per job.

        Args:
            new_summary: LLM-generated summary tailored to job

        Returns:
            True if summary was found and replaced
        """
        # Find summary paragraph (usually first content paragraph)
        summary_para = self._find_summary_paragraph()

        if summary_para:
            original = summary_para.text
            self._set_paragraph_text_preserve_format(summary_para, new_summary)
            self.modifications.append({
                'type': 'summary_rewrite',
                'original': original,
                'modified': new_summary
            })
            return True

        return False
    
    def enhance_bullets(self, bullet_enhancements: List[Dict]) -> List[Dict]:
        """
        Layer 3: Apply suggested bullet improvements.
        These are presented as suggestions for human review.
        
        Args:
            bullet_enhancements: List of {original, suggested, reason}
            
        Returns:
            List of applied enhancements
        """
        applied = []
        
        for enhancement in bullet_enhancements:
            original = enhancement['original']
            suggested = enhancement['suggested']
            
            # Find and replace the bullet
            for paragraph in self.doc.paragraphs:
                if original.strip() in paragraph.text:
                    paragraph.text = paragraph.text.replace(original, suggested)
                    applied.append(enhancement)
                    
                    self.modifications.append({
                        'type': 'bullet_enhancement',
                        'original': original,
                        'modified': suggested,
                        'reason': enhancement.get('reason', '')
                    })
                    break
        
        return applied
    
    def get_summary(self) -> str:
        """Extract current professional summary."""
        summary_para = self._find_summary_paragraph()
        return summary_para.text if summary_para else ""
    
    def get_experience_bullets(self) -> List[str]:
        """Extract experience section bullet points."""
        bullets = []
        in_experience = False
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            
            # Detect experience section
            if re.search(r'(experience|work history|employment)', text, re.IGNORECASE):
                in_experience = True
                continue
            
            # Stop at next major section
            if in_experience and re.search(r'(education|skills|certifications)', text, re.IGNORECASE):
                break
            
            # Collect bullets
            if in_experience and (text.startswith('•') or text.startswith('-')):
                bullets.append(text)
        
        return bullets
    
    def save(self, output_path: str):
        """
        Save modified resume as DOCX.
        
        Args:
            output_path: Path to save tailored resume
        """
        self.doc.save(output_path)
        print(f"✅ Saved tailored resume to: {output_path}")
    
    def get_modification_summary(self) -> Dict:
        """Get summary of all modifications made."""
        return {
            'total_modifications': len(self.modifications),
            'keyword_injections': len([m for m in self.modifications if m['type'] == 'keyword_injection']),
            'summary_rewrites': len([m for m in self.modifications if m['type'] == 'summary_rewrite']),
            'bullet_enhancements': len([m for m in self.modifications if m['type'] == 'bullet_enhancement']),
            'details': self.modifications
        }
    
    def _find_summary_paragraph(self):
        """Find the professional summary paragraph."""
        # Look for summary section header
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.lower()
            if any(keyword in text for keyword in ['summary', 'profile', 'objective', 'about']):
                # Return next paragraph (the actual summary)
                if i + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[i + 1]
        
        # Fallback: assume first paragraph after name is summary
        if len(self.doc.paragraphs) > 2:
            return self.doc.paragraphs[2]
        
        return None
    
    def _add_skills_to_section(self, skills: List[str]) -> int:
        """Add missing skills to Skills section."""
        skills_para = None

        # Find skills section content paragraph (the line after the "Skills" header)
        for i, paragraph in enumerate(self.doc.paragraphs):
            if re.search(r'\bskills\b', paragraph.text, re.IGNORECASE):
                # Use the next paragraph if it has content, otherwise this one
                if i + 1 < len(self.doc.paragraphs) and self.doc.paragraphs[i + 1].text.strip():
                    skills_para = self.doc.paragraphs[i + 1]
                else:
                    skills_para = paragraph
                break

        if not skills_para:
            return 0

        # Get existing skills (word-level check)
        existing_text = skills_para.text.lower()

        # Add missing skills by appending to the last run
        added = 0
        for skill in skills:
            if skill.lower() not in existing_text:
                # Append to existing last run to preserve formatting
                if skills_para.runs:
                    skills_para.runs[-1].text += f", {skill}"
                else:
                    skills_para.add_run(f", {skill}")
                existing_text += f", {skill.lower()}"
                added += 1

        return added
    
    def _replace_in_runs(self, paragraph, pattern: str, replacement: str) -> bool:
        """
        Replace regex pattern within individual runs, preserving run-level formatting.
        Returns True if a replacement was made.
        """
        for run in paragraph.runs:
            if re.search(pattern, run.text, re.IGNORECASE):
                run.text = re.sub(pattern, replacement, run.text, count=1, flags=re.IGNORECASE)
                return True
        return False

    def _set_paragraph_text_preserve_format(self, paragraph, new_text: str):
        """
        Replace all text in a paragraph with new_text while preserving the
        formatting (font, bold, italic, size, color) from the first run.
        Uses direct XML manipulation to avoid python-docx formatting loss.
        """
        # Collect formatting from the first run element
        existing_runs = paragraph._p.findall(qn('w:r'))
        first_rpr = None
        if existing_runs:
            rpr_elem = existing_runs[0].find(qn('w:rPr'))
            if rpr_elem is not None:
                first_rpr = deepcopy(rpr_elem)

        # Remove every run from the paragraph
        for r in existing_runs:
            paragraph._p.remove(r)

        # Build a new run with the same rPr and new text
        new_r = OxmlElement('w:r')
        if first_rpr is not None:
            new_r.append(first_rpr)
        new_t = OxmlElement('w:t')
        new_t.text = new_text
        # Preserve leading/trailing whitespace
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(new_t)
        paragraph._p.append(new_r)

    def _build_synonym_map(self, keywords: List[str]) -> Dict[str, List[str]]:
        """
        Build map of JD keywords to common synonyms.
        Used for intelligent keyword injection.
        """
        # Common synonym patterns
        synonym_patterns = {
            'managed': ['led', 'oversaw', 'directed', 'supervised'],
            'developed': ['built', 'created', 'designed', 'implemented'],
            'improved': ['enhanced', 'optimized', 'increased', 'boosted'],
            'collaborated': ['worked with', 'partnered', 'coordinated'],
            'analyzed': ['examined', 'evaluated', 'assessed', 'reviewed'],
            'leadership': ['management', 'team lead', 'supervision'],
            'communication': ['presentation', 'interpersonal', 'collaboration']
        }
        
        # Build reverse map for keywords
        keyword_map = {}
        for keyword in keywords:
            keyword_lower = keyword.lower()
            # Check if keyword has known synonyms
            if keyword_lower in synonym_patterns:
                keyword_map[keyword] = synonym_patterns[keyword_lower]
            else:
                keyword_map[keyword] = []
        
        return keyword_map


if __name__ == "__main__":
    # Test editor
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python resume_editor.py <path_to_resume.docx>")
        sys.exit(1)
    
    resume_path = sys.argv[1]
    editor = SurgicalResumeEditor(resume_path)
    
    print("Testing Surgical Resume Editor...")
    print(f"\nCurrent Summary:\n{editor.get_summary()}")
    print(f"\nExperience Bullets:\n{editor.get_experience_bullets()[:3]}")
    
    # Test keyword injection
    test_keywords = ["Python", "AWS", "Agile", "Leadership"]
    test_skills = ["Docker", "Kubernetes", "CI/CD"]
    
    injected = editor.inject_keywords(test_keywords, test_skills)
    print(f"\n✅ Injected {injected} keywords")
    
    print(f"\nModifications:\n{editor.get_modification_summary()}")
